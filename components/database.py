import os
from datetime import datetime
import sqlite3
import pandas as pd
from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, Boolean, Float, ForeignKey
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, relationship
import io
import streamlit as st
import uuid

# 创建基类
Base = declarative_base()

class InsurancePolicy(Base):
    """保险方案模型"""
    __tablename__ = 'insurance_policies'

    id = Column(Integer, primary_key=True)
    uuid = Column(String(50), unique=True, nullable=False)
    name = Column(String(200), nullable=False)
    description = Column(Text)
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # 关联的条款版本
    clause_versions = relationship("PolicyClauseVersion", back_populates="policy")

class Clause(Base):
    """条款模型"""
    __tablename__ = 'clauses'

    id = Column(Integer, primary_key=True)
    uuid = Column(String(50), unique=True, nullable=False)
    title = Column(String(200), nullable=False)
    content = Column(Text, nullable=False)
    pinyin = Column(String(200))
    quanpin = Column(String(500))
    insurance_type = Column(String(50))
    company = Column(String(100))
    version = Column(String(20))
    version_number = Column(Integer, default=1)
    is_active = Column(Boolean, default=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # 关联的版本历史
    versions = relationship("ClauseVersion", back_populates="clause")

class ClauseVersion(Base):
    """条款版本历史"""
    __tablename__ = 'clause_versions'

    id = Column(Integer, primary_key=True)
    clause_uuid = Column(String(50), ForeignKey('clauses.uuid'), nullable=False)
    version_number = Column(Integer, nullable=False)
    title = Column(String(200), nullable=False)
    content = Column(Text, nullable=False)
    note = Column(Text)
    created_at = Column(DateTime, default=datetime.utcnow)
    
    # 关联的条款
    clause = relationship("Clause", back_populates="versions")
    # 关联的保险方案版本
    policy_versions = relationship("PolicyClauseVersion", back_populates="clause_version")

class PolicyClauseVersion(Base):
    """保险方案条款版本表"""
    __tablename__ = 'policy_clause_versions'

    id = Column(Integer, primary_key=True)
    policy_id = Column(Integer, ForeignKey('insurance_policies.id'), nullable=False)
    clause_version_id = Column(Integer, ForeignKey('clause_versions.id'), nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    
    # 关联的保险方案
    policy = relationship("InsurancePolicy", back_populates="clause_versions")
    # 关联的条款版本
    clause_version = relationship("ClauseVersion", back_populates="policy_versions")

class Database:
    def __init__(self, db_path=None):
        """初始化数据库连接"""
        # 如果没有提供数据库路径，使用session state中的路径
        if db_path is None and 'db_path' in st.session_state:
            db_path = st.session_state.db_path
        elif db_path is None:
            db_path = 'clauses.db'
        
        # 确保数据库目录存在
        db_dir = os.path.dirname(os.path.abspath(db_path))
        if not os.path.exists(db_dir):
            os.makedirs(db_dir)
        
        # 确保数据库文件所在目录有写权限
        try:
            # 尝试创建一个临时文件来测试写权
            test_file = os.path.join(db_dir, '.test')
            with open(test_file, 'w') as f:
                f.write('test')
            os.remove(test_file)
        except Exception as e:
            st.error(f"数库���录没有写权限: {str(e)}")
            raise
        
        self.db_path = db_path
        self.engine = create_engine(f'sqlite:///{db_path}')
        Base.metadata.create_all(self.engine)
        Session = sessionmaker(bind=self.engine)
        self.session = Session()
        
        # 导出类型供外部使用
        self.Clause = Clause
        self.ClauseVersion = ClauseVersion
        self.InsurancePolicy = InsurancePolicy
        self.PolicyClauseVersion = PolicyClauseVersion

    def create_policy(self, name, description=""):
        """创建保险方案"""
        policy = InsurancePolicy(
            uuid=str(uuid.uuid4()),
            name=name,
            description=description
        )
        self.session.add(policy)
        self.session.commit()
        return policy

    def get_policy(self, policy_id):
        """获取保险方案"""
        return self.session.query(InsurancePolicy).filter_by(id=policy_id).first()

    def update_policy(self, policy_id, name=None, description=None):
        """更新保险方案"""
        policy = self.get_policy(policy_id)
        if policy:
            if name:
                policy.name = name
            if description:
                policy.description = description
            self.session.commit()
            return True
        return False

    def delete_policy(self, policy_id):
        """删除保险方案"""
        policy = self.get_policy(policy_id)
        if policy:
            self.session.delete(policy)
            self.session.commit()
            return True
        return False

    def add_clause_to_policy(self, policy_id, clause_version_id):
        """添加条款版本到保险方案"""
        policy_clause = PolicyClauseVersion(
            policy_id=policy_id,
            clause_version_id=clause_version_id
        )
        self.session.add(policy_clause)
        self.session.commit()
        return policy_clause

    def remove_clause_from_policy(self, policy_id, clause_version_id):
        """从保险方案中移除条款版本"""
        policy_clause = self.session.query(PolicyClauseVersion).filter_by(
            policy_id=policy_id,
            clause_version_id=clause_version_id
        ).first()
        if policy_clause:
            self.session.delete(policy_clause)
            self.session.commit()
            return True
        return False

    def get_policy_clauses(self, policy_id):
        """获取保险方案的所有条款版"""
        return self.session.query(PolicyClauseVersion).filter_by(
            policy_id=policy_id
        ).all()

    def import_clauses(self, df):
        """导入条款数据"""
        # 确保DataFrame包含所需的列
        required_columns = ['UUID', '扩展条款标题', '扩展条款正文', 'PINYIN', 'QUANPIN', '险种', '保险公司', '年度版本']
        if not all(col in df.columns for col in required_columns):
            raise ValueError("导入的数据缺少必要的列")

        # 验证UUID不为空且不重复
        if df['UUID'].isnull().any():
            raise ValueError("存在空的UUID")
        if df['UUID'].duplicated().any():
            raise ValueError("存在重复的UUID")

        # 记录新增和更新的条款数量
        new_count = 0
        update_count = 0

        for _, row in df.iterrows():
            uuid = row['UUID']
            existing_clause = self.session.query(Clause).filter_by(uuid=uuid).first()
            
            if existing_clause:
                # 获取最新版本
                latest_version = self.session.query(ClauseVersion).filter_by(
                    clause_uuid=uuid
                ).order_by(ClauseVersion.version_number.desc()).first()
                
                # 只有当内容有变化时才更新款
                if not latest_version or latest_version.content != row['扩展条款正文']:
                    # 更新条款基本信息，但不创建新版本
                    existing_clause.title = row['扩展条款标题']
                    existing_clause.content = row['扩展条款正文']
                    existing_clause.pinyin = row['PINYIN']
                    existing_clause.quanpin = row['QUANPIN']
                    existing_clause.insurance_type = row['险种']
                    existing_clause.company = row['保险公司']
                    existing_clause.version = row['年度版本']
                    existing_clause.updated_at = datetime.utcnow()
                    
                    # 如果没有任何版本记录，创建初始版本
                    if not latest_version:
                        version = ClauseVersion(
                            clause_uuid=existing_clause.uuid,
                            version_number=1,
                            title=row['扩展条款标题'],
                            content=row['扩展条款正文'],
                            note="初始版本",
                            created_at=datetime.utcnow()
                        )
                        self.session.add(version)
                        existing_clause.version_number = 1
                    
                    update_count += 1
            else:
                # 创建新条款
                new_clause = Clause(
                    uuid=uuid,
                    title=row['扩展条款标题'],
                    content=row['扩展条款正文'],
                    pinyin=row['PINYIN'],
                    quanpin=row['QUANPIN'],
                    insurance_type=row['险种'],
                    company=row['保险公司'],
                    version=row['年度版本'],
                    version_number=1
                )
                self.session.add(new_clause)
                
                # 创建初始版本
                initial_version = ClauseVersion(
                    clause_uuid=uuid,
                    version_number=1,
                    title=row['扩展条款标题'],
                    content=row['扩展条款正文'],
                    note="初始导入",
                    created_at=datetime.utcnow()
                )
                self.session.add(initial_version)
                new_count += 1

            self.session.commit()
        return new_count, update_count

    def export_clauses(self, format='dataframe'):
        """导出条款数据"""
        clauses = self.session.query(Clause).filter_by(is_active=True).all()
        data = []
        for i, clause in enumerate(clauses, 1):
            # 获取最新版本
            latest_version = self.session.query(ClauseVersion).filter_by(
                clause_uuid=clause.uuid
            ).order_by(ClauseVersion.version_number.desc()).first()
            
            # 使用最新版本的内容
            if latest_version:
                content = latest_version.content
                title = latest_version.title
                version_number = latest_version.version_number
            else:
                content = clause.content
                title = clause.title
                version_number = clause.version_number
            
            data.append({
                'UUID': clause.uuid,
                '序号': i,
                '扩展条款标题': title,
                '扩展条款正文': content,
                'PINYIN': clause.pinyin,
                'QUANPIN': clause.quanpin,
                '险种': clause.insurance_type,
                '保险公司': clause.company,
                '年度版本': clause.version,
                '版本号': version_number
            })
        df = pd.DataFrame(data)
        
        if format == 'xlsx':
            output = io.BytesIO()
            df.to_excel(output, index=False)
            output.seek(0)
            return output
        elif format == 'json':
            return df.to_json(orient='records', force_ascii=False)
        elif format == 'dataframe':
            return df
        else:
            return df

    def export_selected_clauses(self, clause_uuids, format='docx'):
        """导出选中的条款"""
        clauses = []
        for uuid in clause_uuids:
            # 获取条款基本信息
            clause = self.session.query(Clause).filter_by(uuid=uuid, is_active=True).first()
            if clause:
                # 获取最新版本
                latest_version = self.session.query(ClauseVersion).filter_by(
                    clause_uuid=uuid
                ).order_by(ClauseVersion.version_number.desc()).first()
                
                if latest_version:
                    # 使用最新版本的内容
                    clause.content = latest_version.content
                    clause.title = latest_version.title
                clauses.append(clause)
        
        if format == 'xlsx':
            data = []
            for i, clause in enumerate(clauses, 1):
                data.append({
                    '序号': i,
                    '扩展条款标题': clause.title,
                    '扩展条款正文': clause.content,
                    '险种': clause.insurance_type,
                    '保险公司': clause.company,
                    '年度版本': clause.version
                })
            df = pd.DataFrame(data)
            output = io.BytesIO()
            df.to_excel(output, index=False)
            output.seek(0)
            return output
        
        elif format == 'docx':
            from docx import Document
            doc = Document()
            for i, clause in enumerate(clauses, 1):
                doc.add_heading(f"{i}. {clause.title}", level=1)
                doc.add_paragraph(clause.content)
                doc.add_paragraph(f"险种：{clause.insurance_type}")
                doc.add_paragraph(f"保险公司：{clause.company}")
                doc.add_paragraph(f"版本：{clause.version}")
                doc.add_page_break()
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output
        
        elif format == 'markdown':
            content = ""
            for i, clause in enumerate(clauses, 1):
                content += f"# {i}. {clause.title}\n\n"
                content += f"{clause.content}\n\n"
                content += f"险种：{clause.insurance_type}\n\n"
                content += f"保险公司：{clause.company}\n\n"
                content += f"版本：{clause.version}\n\n"
                content += "---\n\n"
            return content

    def update_clause(self, uuid, title=None, content=None, version_note=None):
        """更新条款内容"""
        clause = self.session.query(Clause).filter_by(uuid=uuid).first()
        if clause:
            # 获取最新版本号
            latest_version = self.session.query(ClauseVersion).filter_by(
                clause_uuid=uuid
            ).order_by(ClauseVersion.version_number.desc()).first()
            
            # 检查内容是否真的有变化
            if latest_version and latest_version.content == content:
                return True  # 内容没有变化，不需要创建新版本
            
            new_version_number = (latest_version.version_number + 1) if latest_version else 1
            
            # 创建新版本记录
            version = ClauseVersion(
                clause_uuid=clause.uuid,
                version_number=new_version_number,
                title=title if title else clause.title,
                content=content,
                note=version_note if version_note else "手动更新",
                created_at=datetime.utcnow()
            )
            self.session.add(version)
            
            # 更新条款基本信息
            if title:
                clause.title = title
            clause.content = content
            clause.version_number = new_version_number  # 自动切换到新版本
            clause.updated_at = version.created_at
            
            # 提交更改
            self.session.commit()
            
            # 更新session state中的条款内容
            if 'selected_clauses' in st.session_state:
                st.session_state.selected_clauses = [
                    {
                        **c,
                        '扩展���款正文': content,
                        '版本号': new_version_number,
                        '扩展条款标题': title if title else c['扩展条款标题']
                    } if c['UUID'] == uuid else c
                    for c in st.session_state.selected_clauses
                ]
            
            return True
        return False

    def get_clause_versions(self, uuid):
        """获取条款的所有版本"""
        # 获取当前版本
        current = self.session.query(Clause).filter_by(uuid=uuid).first()
        if not current:
            return []
        
        # 获取所有版本（包括当前版本）
        versions = self.session.query(ClauseVersion).filter_by(
            clause_uuid=uuid
        ).order_by(ClauseVersion.version_number.desc()).all()
        
        # 如果没有版本记录，创建初始版本
        if not versions:
            initial_version = ClauseVersion(
                clause_uuid=current.uuid,
                version_number=1,
                title=current.title,
                content=current.content,
                note="初始版本",
                created_at=current.created_at
            )
            self.session.add(initial_version)
            self.session.commit()
            versions = [initial_version]
        
        return versions

    def activate_clause_version(self, uuid, version_number):
        """激活指定版本的条款"""
        print(f"DEBUG: 开始激活版本 - UUID: {uuid}, 版本号: {version_number}")
        
        # 获取条款
        clause = self.session.query(Clause).filter_by(uuid=uuid).first()
        if not clause:
            print("DEBUG: 未找到条款")
            return False
        
        print(f"DEBUG: 当前条款版本号: {clause.version_number}")
        
        # 获取要激活的版本
        version = self.session.query(ClauseVersion).filter_by(
            clause_uuid=uuid,
            version_number=version_number
        ).first()
        
        if version:
            try:
                print(f"DEBUG: 找到目标版本 - 标题: {version.title}")
                print(f"DEBUG: 目标版本内容: {version.content[:50]}...")
                
                # 更新当前条款为选中的版本
                clause.title = version.title
                clause.content = version.content
                clause.version_number = version.version_number
                clause.updated_at = datetime.utcnow()
                
                # 立即提交更改
                self.session.commit()
                print("DEBUG: 数据库更新成功")
                
                # 更新session state中的条款内容
                if 'selected_clauses' in st.session_state:
                    print("DEBUG: 开始更新session state中的条款")
                    old_clauses = len(st.session_state.selected_clauses)
                    st.session_state.selected_clauses = [
                        {
                            **c,
                            '扩展条款正文': version.content,
                            '版本号': version.version_number,
                            '扩展条款标题': version.title
                        } if c['UUID'] == uuid else c
                        for c in st.session_state.selected_clauses
                    ]
                    print(f"DEBUG: 更新后的条款数量: {len(st.session_state.selected_clauses)}")
                    
                    # 强制保存到数据库
                    if 'current_policy_id' in st.session_state:
                        print("DEBUG: 保存更改到保险方案")
                        clause_uuids = [c['UUID'] for c in st.session_state.selected_clauses]
                        self.save_policy_clauses(st.session_state.current_policy_id, clause_uuids)
                
                # 强制刷新session state中的版本信息
                if 'version_info' not in st.session_state:
                    st.session_state.version_info = {}
                st.session_state.version_info[uuid] = version_number
                print(f"DEBUG: 更新version_info: {st.session_state.version_info}")
                
                return True
                
            except Exception as e:
                self.session.rollback()
                print(f"DEBUG: 版本切换失败: {str(e)}")
                return False
        
        print("DEBUG: 未找到目标版本")
        return False

    def delete_clause_version(self, uuid, version_number):
        """删除指定版的条款"""
        # 获取所有版本
        versions = self.get_clause_versions(uuid)
        
        # 如果只有一个版本，不允许删除
        if len(versions) <= 1:
            return False
        
        # 如果要删除的是当前版本，不允许删除
        clause = self.session.query(Clause).filter_by(uuid=uuid).first()
        if clause.version_number == version_number:
            return False
        
        # 删除指定版本
        self.session.query(ClauseVersion).filter_by(
            clause_uuid=uuid,
            version_number=version_number
        ).delete()
        
        self.session.commit()
        return True

    def clear_database(self):
        """清空数据库"""
        self.session.query(Clause).delete()
        self.session.query(ClauseVersion).delete()
        self.session.query(InsurancePolicy).delete()
        self.session.query(PolicyClauseVersion).delete()
        self.session.commit()

    def export_database(self):
        """导出数据库"""
        if os.path.exists(self.db_path):
            with open(self.db_path, 'rb') as f:
                return f.read()
        return None

    def import_database(self, db_data):
        """导入数据库"""
        try:
            # 先关闭当前会话
            self.session.close()
            
            # 备份现有数据库（如果存在）
            if os.path.exists(self.db_path):
                backup_path = f"{self.db_path}.bak"
                os.rename(self.db_path, backup_path)
            
            # 写入新数据库
            with open(self.db_path, 'wb') as f:
                f.write(db_data)
            
            # 重新创建会话
            self.engine = create_engine(f'sqlite:///{self.db_path}')
            Session = sessionmaker(bind=self.engine)
            self.session = Session()
            
            return True
        except Exception as e:
            print(f"导入数据库失败：{str(e)}")
            # 恢复备份
            if os.path.exists(f"{self.db_path}.bak"):
                if os.path.exists(self.db_path):
                    os.remove(self.db_path)
                os.rename(f"{self.db_path}.bak", self.db_path)
            return False
        finally:
            # 清理备份文件
            if os.path.exists(f"{self.db_path}.bak"):
                os.remove(f"{self.db_path}.bak")

    def __del__(self):
        """关闭数据库连接"""
        self.session.close()

    def get_policy_by_uuid(self, uuid):
        """通过 UUID 获取保险方案"""
        return self.session.query(InsurancePolicy).filter_by(uuid=uuid).first()

    def get_clause_version_by_clause_uuid(self, clause_uuid):
        """获取条款的最新版本"""
        return self.session.query(ClauseVersion).filter_by(
            clause_uuid=clause_uuid
        ).order_by(ClauseVersion.version_number.desc()).first()

    def save_policy_clauses(self, policy_id, clause_uuids):
        """保存保险方案关联的条款"""
        print(f"保存条款到保险方案，保险方案ID：{policy_id}，条款数量：{len(clause_uuids)}")
        
        try:
            # 获取现有的关联记录
            existing_relations = self.session.query(PolicyClauseVersion).filter_by(
                policy_id=policy_id
            ).all()
            
            # 创建映射：clause_uuid -> clause_version_id
            existing_versions = {}
            for relation in existing_relations:
                clause_version = relation.clause_version
                if clause_version:
                    existing_versions[clause_version.clause_uuid] = relation
            
            print(f"现有关联数量：{len(existing_versions)}")
            
            # 处理每个条款
            updated_count = 0
            for clause_uuid in clause_uuids:
                # 获取条款的当前版本号
                current_version_number = None
                if 'version_info' in st.session_state:
                    current_version_number = st.session_state.version_info.get(clause_uuid)
                
                # 如果session state中没有版本信息，从数据库获取
                if current_version_number is None:
                    clause = self.session.query(Clause).filter_by(uuid=clause_uuid).first()
                    if clause:
                        current_version_number = clause.version_number
                
                print(f"处理条款 {clause_uuid}，版本号：{current_version_number}")
                
                # 获取指定版本的条款版本记录
                clause_version = self.session.query(ClauseVersion).filter_by(
                    clause_uuid=clause_uuid,
                    version_number=current_version_number
                ).first()
                
                if clause_version:
                    if clause_uuid in existing_versions:
                        # 更新现有关联
                        relation = existing_versions[clause_uuid]
                        if relation.clause_version_id != clause_version.id:
                            relation.clause_version_id = clause_version.id
                            updated_count += 1
                    else:
                        # 创建新关联
                        new_relation = PolicyClauseVersion(
                            policy_id=policy_id,
                            clause_version_id=clause_version.id
                        )
                        self.session.add(new_relation)
                        updated_count += 1
            
            # 删除不再需要的关联
            for clause_uuid, relation in existing_versions.items():
                if clause_uuid not in clause_uuids:
                    self.session.delete(relation)
                    updated_count += 1
            
            if updated_count > 0:
                self.session.commit()
                print(f"更新了 {updated_count} 个关联")
            
            return True
            
        except Exception as e:
            self.session.rollback()
            print(f"保存条款关联失败：{str(e)}")
            return False

    def get_policy_clause_uuids(self, policy_id):
        """获取保险方案关联的所有条款UUID"""
        print(f"获取保险方案条款，保险方案ID：{policy_id}")
        policy_clauses = self.session.query(PolicyClauseVersion).filter_by(
            policy_id=policy_id
        ).join(ClauseVersion).all()
        
        uuids = [pc.clause_version.clause_uuid for pc in policy_clauses]
        print(f"找到关联条款数量：{len(uuids)}")
        return uuids
