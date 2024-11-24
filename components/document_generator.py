import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io

def create_element(name):
    """创建XML元素"""
    return OxmlElement(name)

def create_attribute(element, name, value):
    """创建XML属性"""
    element.set(qn(name), value)

def add_bookmark(paragraph, bookmark_name):
    """添加书签"""
    run = paragraph.add_run()
    tag = create_element('w:bookmarkStart')
    create_attribute(tag, 'w:id', '0')
    create_attribute(tag, 'w:name', bookmark_name)
    run._r.append(tag)
    
    tag = create_element('w:bookmarkEnd')
    create_attribute(tag, 'w:id', '0')
    run._r.append(tag)
    
    return run

def add_hyperlink(paragraph, text, bookmark_name):
    """添加超链接"""
    run = paragraph.add_run()
    hyperlink = create_element('w:hyperlink')
    create_attribute(hyperlink, 'w:anchor', bookmark_name)
    
    run = create_element('w:r')
    rPr = create_element('w:rPr')
    
    # 添加超链接样式
    rStyle = create_element('w:rStyle')
    create_attribute(rStyle, 'w:val', 'Hyperlink')
    rPr.append(rStyle)
    
    run.append(rPr)
    text_element = create_element('w:t')
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    
    paragraph._p.append(hyperlink)

def generate_markdown(insurance_data, selected_clauses):
    """生成带目录和跳转的Markdown格式保险方案"""
    markdown = f"""# 保险方案

# 投保人
名称：{insurance_data['policyholder']}

# 被保险人
名称：{insurance_data['insured']['name']}
证件类型：{insurance_data['insured']['id_type']} 证件号码：{insurance_data['insured']['id_number']}
联系人：{insurance_data['insured']['contact']['name']} 联系人电话：{insurance_data['insured']['contact']['phone']} 联系人邮箱：{insurance_data['insured']['contact']['email']}
联系地址：{insurance_data['insured']['contact']['address']} 联系邮编：{insurance_data['insured']['contact']['postal_code']}

# 被保险人标的地址
名称：{insurance_data['property']['name']}
地址：{insurance_data['property']['address']}

# 主险
## 第一部分 物质损失
"""
    
    # 添加物质损失表格
    markdown += "\n| 标的类别 | 保险金额（元） | 费率（%） | 保费（元） |\n"
    markdown += "|----------|--------------|---------|----------|\n"
    for item in insurance_data['material_loss']:
        markdown += f"| {item['标的类别']} | {item['保险金额（元）']} | {item['费率（%）']} | {item['保费（元）']} |\n"
    
    # 添加第三者责任表格
    markdown += "\n## 第二部分 第三者责任\n"
    markdown += "\n| 限额名称 | 责任限额（元） | 保费（元） |\n"
    markdown += "|----------|--------------|----------|\n"
    for item in insurance_data['liability']:
        markdown += f"| {item['限额名称']} | {item['责任限额（元）']} | {item['保费（元）']} |\n"
    
    # 添加免赔额表格
    markdown += "\n## 第三部分 免赔额\n"
    markdown += "\n| 免赔项目 | 免赔额 / 免赔约定 |\n"
    markdown += "|----------|------------------|\n"
    for item in insurance_data['deductibles']:
        markdown += f"| {item['免赔项目']} | {item['免赔额 / 免赔约定']} |\n"
    
    # 添加其他信息部分
    if 'other_info_tabs' in insurance_data and 'other_info_data' in insurance_data:
        markdown += "\n# 其他信息\n"
        
        for tab in insurance_data['other_info_tabs']:
            tab_data = insurance_data['other_info_data'].get(tab['id'], [])
            if tab_data:
                markdown += f"\n## {tab['name']}\n"
                markdown += "\n| 项目 | 内容说明 |\n"
                markdown += "|------|----------|\n"
                for item in tab_data:
                    markdown += f"| {item['项目']} | {item['内容说明']} |\n"
    
    # 添加特别约定
    if 'special_terms' in insurance_data and insurance_data['special_terms']:
        markdown += "\n# 特别约定\n\n"
        for i, term in enumerate(insurance_data['special_terms'], 1):
            markdown += f"{i}. {term}\n\n"
    
    # 生成目录锚点
    toc_links = []
    for i, clause in enumerate(selected_clauses, 1):
        # 创建锚点ID
        anchor_id = f"clause-{i}"
        # 添加目录项
        toc_links.append(f"{i}. [{clause['扩展条款标题']}](#{anchor_id})")
    
    # 添加扩展条款目录
    markdown += "\n# 扩展条款目录\n\n"
    markdown += "\n".join(toc_links)
    
    # 添加扩展条款内容（带锚点）
    markdown += "\n\n# 扩展条款\n\n"
    for i, clause in enumerate(selected_clauses, 1):
        markdown += f"<a id='clause-{i}'></a>\n\n"
        markdown += f"## {i}. {clause['扩展条款标题']}\n\n"
        markdown += f"{clause['扩展条款正文']}\n\n"
    
    return markdown

def generate_docx(insurance_data, selected_clauses):
    """生成带目录的DOCX格式保险方案"""
    doc = Document()
    
    # 设置默认字体为仿宋
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # 设置一级标题样式（文档标题）
    title_style = doc.styles['Heading 1']
    title_style.font.name = '仿宋'
    title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    
    # 设置二级标题样式（主要章节：投保人、被保险人等）
    heading1_style = doc.styles['Heading 2']
    heading1_style.font.name = '黑体'
    heading1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading1_style.font.size = Pt(14)
    heading1_style.font.bold = True
    
    # 设置三级标题样式（子章节：物质损失、第三者责任等）
    heading2_style = doc.styles['Heading 3']
    heading2_style.font.name = '黑体'
    heading2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading2_style.font.size = Pt(12)
    heading2_style.font.bold = True
    
    # 添加文档标题
    doc.add_heading('保险方案', 0)
    
    # 投保人信息
    doc.add_heading('投保人', level=2)  # 改为二级标题
    doc.add_paragraph(f"名称：{insurance_data['policyholder']}")
    
    # 被保险人信息
    doc.add_heading('被保险人', level=2)  # 改为二级标题
    doc.add_paragraph(f"名称：{insurance_data['insured']['name']}")
    doc.add_paragraph(f"证件类型：{insurance_data['insured']['id_type']} 证件号码：{insurance_data['insured']['id_number']}")
    doc.add_paragraph(
        f"联系人：{insurance_data['insured']['contact']['name']} "
        f"联系人电话：{insurance_data['insured']['contact']['phone']} "
        f"联系人邮箱：{insurance_data['insured']['contact']['email']}"
    )
    doc.add_paragraph(
        f"联系地址：{insurance_data['insured']['contact']['address']} "
        f"联系邮编：{insurance_data['insured']['contact']['postal_code']}"
    )
    
    # 被保险人标的地址
    doc.add_heading('被保险人标的地址', level=2)  # 改为二级标题
    doc.add_paragraph(f"名称：{insurance_data['property']['name']}")
    doc.add_paragraph(f"地址：{insurance_data['property']['address']}")
    
    # 主险
    doc.add_heading('主险', level=2)  # 改为二级标题
    
    # 物质损失表格
    doc.add_heading('第一部分 物质损失', level=3)  # 使用三级标题
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = '标的类别'
    header_cells[1].text = '保险金额（元）'
    header_cells[2].text = '费率（%）'
    header_cells[3].text = '保费（元）'
    
    for item in insurance_data['material_loss']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['标的类别'])
        row_cells[1].text = str(item['保险金额（元）'])
        row_cells[2].text = str(item['费率（%）'])
        row_cells[3].text = str(item['保费（元）'])
    
    # 第三者责任表格
    doc.add_heading('第二部分 第三者责任', level=3)  # 使用三级标题
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = '限额名称'
    header_cells[1].text = '责任限额（元）'
    header_cells[2].text = '保费（元）'
    
    for item in insurance_data['liability']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['限额名称'])
        row_cells[1].text = str(item['责任限额（元）'])
        row_cells[2].text = str(item['保费（元）'])
    
    # 免赔额表格
    doc.add_heading('第三部分 免赔额', level=3)  # 使用三级标题
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = '免赔项目'
    header_cells[1].text = '免赔额 / 免赔约定'
    
    for item in insurance_data['deductibles']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item['免赔项目'])
        row_cells[1].text = str(item['免赔额 / 免赔约定'])
    
    # 添加其他信息部分
    if 'other_info_tabs' in insurance_data and 'other_info_data' in insurance_data:
        doc.add_heading('其他信息', level=2)  # 改为二级标题
        
        for tab in insurance_data['other_info_tabs']:
            tab_data = insurance_data['other_info_data'].get(tab['id'], [])
            if tab_data:
                doc.add_heading(tab['name'], level=3)  # 使用三级标题
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                header_cells = table.rows[0].cells
                header_cells[0].text = '项目'
                header_cells[1].text = '内容说明'
                
                for item in tab_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item['项目'])
                    row_cells[1].text = str(item['内容说明'])
    
    # 特别约定
    if 'special_terms' in insurance_data and insurance_data['special_terms']:
        doc.add_heading('特别约定', level=2)  # 改为二级标题
        for i, term in enumerate(insurance_data['special_terms'], 1):
            doc.add_paragraph(f"{i}. {term}")
    
    # 扩展条款目录
    doc.add_heading('扩展条款目录', level=2)  # 改为二级标题
    for i, clause in enumerate(selected_clauses, 1):
        paragraph = doc.add_paragraph()
        # 添加目录项和超链接
        add_hyperlink(paragraph, f"{i}. {clause['扩展条款标题']}", f"clause_{i}")
    
    # 添加分页符
    doc.add_page_break()
    
    # 扩展条款
    doc.add_heading('扩展条款', level=2)  # 改为二级标题
    for i, clause in enumerate(selected_clauses, 1):
        # 添加条款标题（带书签）
        paragraph = doc.add_paragraph()
        add_bookmark(paragraph, f"clause_{i}")
        paragraph.add_run(f"{i}. {clause['扩展条款标题']}")
        paragraph.style = 'Heading 3'  # 使用三级标题
        
        # 添加条款内容（移除换行符）
        content = clause['扩展条款正文'].replace('\n', ' ').strip()
        doc.add_paragraph(content)
    
    # 保存文档到内存
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer

def generate_document(insurance_data, selected_clauses, format='markdown'):
    """根据指定格式生成文档"""
    if format.lower() == 'markdown':
        return generate_markdown(insurance_data, selected_clauses)
    elif format.lower() == 'docx':
        return generate_docx(insurance_data, selected_clauses)
    else:
        raise ValueError(f"Unsupported format: {format}")
